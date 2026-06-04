from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "铁路投资控制插件功能说明参考文档.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=3)
    run = p.add_run(title)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)
    run.bold = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=12)
    run = p.add_run(subtitle)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before=0, after=3, line=1.05)
    return table


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("铁路投资控制插件功能说明")


def build_doc():
    doc = Document()
    apply_styles(doc)

    today = date.today().strftime("%Y-%m-%d")
    add_title(
        doc,
        "铁路投资控制插件功能说明参考文档",
        f"参考《Excel 工程数量表智能定额映射系统 - 用户操作手册》和《工程造价智能编制及分析系统 - 详细设计说明书》整理；生成日期：{today}",
    )

    add_h1(doc, "1. 系统定位")
    add_body(
        doc,
        "本插件体系面向铁路基本建设工程投资控制系统，目标是在不改变原软件主流程的前提下，补充工程量处理、定额推荐、Excel 联动和人工经验沉淀能力。与独立 Web 上传式系统不同，本方案直接嵌入既有造价软件界面，围绕现场正在编辑的定额输入表、章节树和 Excel/WPS 工程数量表展开操作。",
    )
    add_bullets(
        doc,
        [
            "低侵入：保留原软件计算、录入、复制粘贴和数据库逻辑，插件只在必要入口辅助操作。",
            "现场联动：优先读取当前 Excel/WPS 选区、当前定额输入表和当前章节，减少文件上传和任务切换。",
            "经验复用：通过 mapping-boxes.jsonl 保存人工扶正后的工程量表达与一组定额/材料关系。",
            "人控优先：不可靠推荐允许为空，等待人工扶正；组件框和 AI 结果均以用户确认为准。",
        ],
    )

    add_h1(doc, "2. 功能总览")
    add_table(
        doc,
        ["功能", "入口/对象", "主要输入", "输出结果", "定位"],
        [
            ["乘系数", "定额输入表或左侧章节树右键菜单", "选中定额行、当前章节、系数表达", "批量更新工程数量输入或定额编号表达式", "现场批量调整工具"],
            ["Excel 工程量联动", "定额输入表右键菜单、快捷键、联动面板", "定额行、Excel 单元格或表达式", "保存绑定关系并同步工程数量", "持续跟踪 Excel 数量变化"],
            ["推荐定额", "定额输入表右键“推荐定额”", "Excel/WPS 当前选区或剪贴板工程量", "推荐定额/组件框，复制为可粘贴内容", "批量从工程量生成定额输入候选"],
            ["添加 AI 配对框内容", "定额输入表右键“添加对应框内容”", "选中定额/材料与 Excel 工程量选区", "写入 mapping-boxes.jsonl 的组件框", "人工经验快速沉淀"],
        ],
        [1700, 2200, 2200, 1800, 1460],
    )

    add_h1(doc, "3. 乘系数功能")
    add_h2(doc, "3.1 使用场景")
    add_body(
        doc,
        "当一组定额需要统一乘以或除以某个系数时，用户可以在定额输入表选中若干行，或在左侧章节树选择一个章节，然后通过右键菜单执行乘系数。该功能更接近原软件内的批量编辑增强，而不是智能映射系统里的单位换算。",
    )
    add_h2(doc, "3.2 工作流程")
    add_numbered(
        doc,
        [
            "用户在定额输入表选中需要调整的定额行，或在左侧章节树选中章节。",
            "右键点击“乘系数”，选择“乘到原来的工程量”或“乘到定额编号”。",
            "输入系数，例如 *0.8、*1.05 或 /1.1。",
            "插件更新定额输入表中的工程数量输入、工程数量，或把系数追加到定额编号表达式。",
            "刷新当前定额输入表，提示已处理的定额条数。",
        ],
    )
    add_h2(doc, "3.3 与参考 PDF 中系数机制的差别")
    add_body(
        doc,
        "参考 PDF 中的系数主要用于单位识别和数量换算，例如识别 100m2、100m3 后计算标准定额数量。本插件的乘系数是明确的人工批量调整动作，作用对象是已经录入软件的定额行或章节下定额。两者可以互补：推荐定额负责把 Excel 数量换算成定额数量，乘系数负责后续人工批量修正。",
    )

    add_h1(doc, "4. Excel 工程量联动")
    add_h2(doc, "4.1 使用场景")
    add_body(
        doc,
        "Excel 联动用于解决工程数量表后续变化后，软件中定额工程数量需要同步更新的问题。用户可以把某条定额绑定到 Excel 单元格，也可以绑定到由多个单元格组成的表达式，例如 E4+E5 或 E4*0.8。",
    )
    add_h2(doc, "4.2 支持能力")
    add_bullets(
        doc,
        [
            "单条绑定：当前定额行绑定 Excel/WPS 当前单元格。",
            "连续绑定：从 Excel 起始单元格向下，对应绑定多条连续定额行。",
            "表达式绑定：支持求和、乘除和手写公式，保存表达式而不是只保存单元格。",
            "AI 智能匹配：读取软件中选中的定额与 Excel 选区，给出单元格或表达式绑定建议。",
            "同步面板：集中查看、手动同步、删除绑定，并记录最近同步状态。",
        ],
    )
    add_h2(doc, "4.3 数据保存与同步")
    add_body(
        doc,
        "联动关系保存为项目相关的本地 XML 结构，记录项目标识、定额序号、定额编号、Excel 文件路径、工作表、单元格地址、表达式、最近同步值和状态。同步时优先读取正在打开的 Excel/WPS；如果无法通过 COM 读取，则尝试从文件读取。同步成功后更新定额输入表的工程数量输入和工程数量，并写入同步日志。",
    )

    add_h1(doc, "5. 推荐定额功能")
    add_h2(doc, "5.1 使用场景")
    add_body(
        doc,
        "推荐定额用于从 Excel/WPS 工程数量表批量生成定额输入候选。用户在 Excel 中框选工程量名称、单位和数量，回到软件定额输入表右键打开推荐窗口，插件按行识别工程量并给出推荐定额或组件框。",
    )
    add_h2(doc, "5.2 推荐优先级")
    add_numbered(
        doc,
        [
            "先查 mapping-boxes.jsonl 中人工扶正后的组件框。",
            "组件框未命中时查本地 quota-index.jsonl 全量定额索引。",
            "仍不可靠时可由 AI 补推，但候选必须来自本地索引或组件框候选。",
            "没有可靠结果时保留空推荐行，等待人工扶正。",
        ],
    )
    add_h2(doc, "5.3 输出方式")
    add_body(
        doc,
        "插件不直接绕开原软件业务逻辑写项目数据库。用户确认推荐后点击“复制勾选内容”，插件把定额编号和换算后的工程数量放入剪贴板，用户在定额输入表目标位置粘贴。普通检索只返回一条最优定额；材料只能通过人工扶正后的组件框出现。",
    )

    add_h1(doc, "6. 添加 AI 配对框内容")
    add_h2(doc, "6.1 功能定位")
    add_body(
        doc,
        "添加 AI 配对框内容是 mapping-boxes.jsonl 的训练入口。它不是另一个推荐系统，而是把用户现场确认过的“工程量表达 -> 一组定额/材料”沉淀为组件框，让后续推荐优先复用。",
    )
    add_h2(doc, "6.2 工作流程")
    add_numbered(
        doc,
        [
            "用户在定额输入表中框选一条或多条正确的定额/材料。",
            "用户在 Excel/WPS 中框选对应的工程量行。",
            "打开“添加对应框内容”窗口，点击 AI 配对。",
            "系统结合本地数量匹配和 AI 语义判断，生成配对候选。",
            "用户检查、勾选并确认后，写入 mapping-boxes.jsonl。",
        ],
    )
    add_h2(doc, "6.3 与“批量反向建模板”的关系")
    add_body(
        doc,
        "两者目标相同，都是生成组件框。当前功能适合小范围、现场手动训练；批量反向建模板则是把范围扩大到整章或整项目。由于两者本质接近，当前阶段保留现有“添加 AI 配对框内容”即可，后续如果需要再扩展为批量模式。",
    )

    add_h1(doc, "7. 与参考 PDF 系统的主要差异")
    add_table(
        doc,
        ["对比项", "参考 PDF 系统", "本插件体系"],
        [
            ["系统形态", "FastAPI Web 服务，上传 Excel 后异步处理", "嵌入铁路投资控制系统，围绕当前软件界面操作"],
            ["数据底座", "MySQL + Milvus 向量库 + LLM", "本地 JSONL 学习库、定额索引、软件项目库和可选 AI"],
            ["Excel 处理", "按文件上传任务处理", "优先读取当前 Excel/WPS 选区、剪贴板或当前单元格"],
            ["推荐逻辑", "向量检索、模板复制、对比任务", "mapping-boxes 优先、quota-index 检索、AI 补推、人工扶正"],
            ["数量变化处理", "对比模式识别新旧 Excel 差异", "通过 Excel 联动保存绑定并同步变化"],
            ["人工学习", "用户确认映射后写入向量知识库", "人工扶正和 AI 配对写入 mapping-boxes.jsonl"],
            ["写入方式", "生成定额输入数据或写入服务端数据库", "推荐结果通过剪贴板粘贴；乘系数和联动按原软件数据结构更新"],
        ],
        [1700, 3800, 3860],
    )

    add_h1(doc, "8. 数据与安全边界")
    add_bullets(
        doc,
        [
            "mapping-boxes.jsonl 保存人工确认后的组件框，是推荐链路的最高优先级学习库。",
            "quota-index.jsonl 和 material-index.jsonl 是本地索引缓存，普通推荐只返回定额，不直接返回材料。",
            "Excel 联动保存的是项目内绑定关系和同步日志，不应替代原始工程数量表管理。",
            "推荐定额仍采用剪贴板粘贴回原软件的方式，避免绕开原软件业务逻辑。",
            "不删除用户已有学习库、索引库、日志和业务目录数据。",
        ],
    )

    add_h1(doc, "9. 后续扩展建议")
    add_numbered(
        doc,
        [
            "先继续强化现有添加 AI 配对框内容功能，提高组件框确认体验和错误回退能力。",
            "在推荐定额中继续优化特殊工程量规则，例如钢筋、警示带、检查井、土方外运等高频误匹配场景。",
            "Excel 联动可作为设计变更处理的主线能力，后续再考虑新旧 Excel 对比报告。",
            "暂不单独实施批量反向建模板，避免与现有 AI 配对框功能重复。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
