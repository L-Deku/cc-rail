from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "铁路投资控制插件功能操作手册与设计说明.docx"


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def para_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def run_font(run, size=None, color=None, bold=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("铁路投资控制插件功能操作手册与设计说明")
    run_font(r, size=9, color=RGBColor(85, 85, 85))


def title(doc, text, subtitle):
    p = doc.add_paragraph()
    para_spacing(p, after=4)
    r = p.add_run(text)
    run_font(r, size=22, color=RGBColor(31, 77, 120), bold=True)
    p = doc.add_paragraph()
    para_spacing(p, after=12)
    r = p.add_run(subtitle)
    run_font(r, size=10, color=RGBColor(85, 85, 85))


def h1(doc, text):
    return doc.add_paragraph(text, style="Heading 1")


def h2(doc, text):
    return doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    return doc.add_paragraph(text, style="Heading 3")


def body(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def callout(doc, label, text, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + "：")
    run_font(r, bold=True)
    p.add_run(text)
    para_spacing(p, after=2, line=1.05)
    doc.add_paragraph()


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        set_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_geometry(tbl, widths)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                para_spacing(p, after=3, line=1.05)
    doc.add_paragraph()
    return tbl


def build():
    doc = Document()
    apply_styles(doc)

    title(
        doc,
        "铁路投资控制插件功能操作手册与设计说明",
        "参考《Excel 工程数量表智能定额映射系统 - 用户操作手册》和《工程造价智能编制及分析系统 - 详细设计说明书》的结构整理",
    )

    table(
        doc,
        ["项目", "内容"],
        [
            ["文档名称", "铁路投资控制插件功能操作手册与设计说明"],
            ["版本号", "v1.0"],
            ["生成日期", date.today().strftime("%Y-%m-%d")],
            ["适用范围", "乘系数、Excel 工程量联动、推荐定额、添加 AI 配对框内容"],
            ["数据原则", "本地文件优先；人工扶正优先；推荐结果不绕开原软件业务流程"],
        ],
        [1800, 7560],
    )

    h1(doc, "目录")
    numbered(
        doc,
        [
            "概述",
            "系统架构设计",
            "四种处理模式",
            "模式一：乘系数模式",
            "模式二：Excel 工程量联动模式",
            "模式三：推荐定额模式",
            "模式四：AI 配对框学习模式",
            "用户映射学习功能",
            "典型工作场景",
            "常见问题解答",
            "操作技巧和建议",
            "附录：数据文件和术语说明",
        ],
    )

    h1(doc, "1. 概述")
    h2(doc, "1.1 系统简介")
    body(
        doc,
        "本插件体系是在铁路基本建设工程投资控制系统基础上增加的辅助能力。它不是独立造价平台，而是围绕原软件的定额输入表、章节树、Excel/WPS 工程数量表和本地学习库工作。核心目标是减少重复录入、提高工程数量与定额之间的对应效率，并把人工确认过的经验沉淀下来。",
    )
    body(
        doc,
        "参考 PDF 中的系统采用上传 Excel、异步任务、MySQL、Milvus 向量库和大语言模型完成智能映射。本插件的路线更轻：直接嵌入当前软件，读取正在操作的 Excel 或剪贴板，推荐结果通过原软件界面粘贴或同步，学习结果保存在本地 mapping-boxes.jsonl。",
    )

    h2(doc, "1.2 核心价值")
    bullet(
        doc,
        [
            "减少重复操作：批量乘系数、批量绑定 Excel、批量推荐定额，减少逐条录入。",
            "保留原软件逻辑：推荐定额采用剪贴板粘贴，计算和业务校验仍由原软件承担。",
            "提高变更响应速度：Excel 联动可以在工程数量表变化后同步更新定额工程数量。",
            "沉淀人工经验：扶正和 AI 配对写入 mapping-boxes.jsonl，下次优先复用。",
            "允许不确定：不可靠推荐不强塞，保留空行和扶正入口，避免错误扩散。",
        ],
    )

    h2(doc, "1.3 核心功能")
    table(
        doc,
        ["序号", "功能", "一句话说明"],
        [
            ["1", "乘系数", "对选中定额或章节下定额批量乘除系数。"],
            ["2", "Excel 工程量联动", "把定额工程数量绑定到 Excel 单元格或表达式，后续可同步。"],
            ["3", "推荐定额", "从 Excel 工程量表读取名称、单位和数量，推荐定额或组件框。"],
            ["4", "添加 AI 配对框内容", "把人工确认的工程量与定额/材料组合保存为可复用组件框。"],
        ],
        [900, 2300, 6160],
    )

    h1(doc, "2. 系统架构设计")
    h2(doc, "2.1 整体架构")
    body(
        doc,
        "整体架构可以理解为三层：宿主软件界面层、插件功能层、本地数据层。宿主软件界面层提供定额输入表、章节树和项目数据库连接；插件功能层提供乘系数、Excel 联动、推荐定额和 AI 配对；本地数据层保存定额索引、材料索引、组件框和联动配置。",
    )
    table(
        doc,
        ["层级", "组成", "职责"],
        [
            ["宿主软件界面层", "定额输入表、章节树、项目数据库连接、右键菜单", "提供用户正在操作的上下文和原软件业务入口。"],
            ["插件功能层", "RecoExpandPanel、RecoQuotaRecommend", "挂接右键菜单，读取 Excel/WPS，生成推荐，执行联动或批量调整。"],
            ["本地数据层", "quota-index.jsonl、material-index.jsonl、mapping-boxes.jsonl、Excel 联动 XML", "保存索引、学习库、组件框和项目绑定关系。"],
        ],
        [1900, 3200, 4260],
    )

    h2(doc, "2.2 模块划分")
    bullet(
        doc,
        [
            "乘系数模块：负责右键菜单注入、系数输入、选中行或章节范围识别、数据库更新和界面刷新。",
            "Excel 联动模块：负责 Excel/WPS COM 读取、手动绑定、连续绑定、表达式绑定、AI 智能匹配、同步面板和日志记录。",
            "推荐定额模块：负责读取 Excel 选区或剪贴板、工程量行解析、组件框查找、索引检索、AI 补推、复制粘贴内容生成。",
            "AI 配对框模块：负责读取定额输入表选中目标和 Excel 选区工程量，生成配对候选并写入 mapping-boxes。",
        ],
    )

    h2(doc, "2.3 数据流设计")
    body(doc, "推荐定额的数据流：Excel 选区 -> 工程量行解析 -> mapping-boxes 优先匹配 -> quota-index 检索 -> AI 补推 -> 用户勾选 -> 剪贴板粘贴。")
    body(doc, "Excel 联动的数据流：定额行 -> Excel 单元格/表达式 -> 保存绑定 -> 读取 Excel 当前值 -> 更新定额输入表 -> 写同步日志。")
    body(doc, "AI 配对框的数据流：选中定额/材料 -> Excel 工程量选区 -> 本地数量匹配和 AI 语义配对 -> 用户确认 -> 写入 mapping-boxes。")

    h1(doc, "3. 四种处理模式")
    body(doc, "本插件可以按四种处理模式理解。每种模式面向不同工作场景，入口不同，但都围绕当前软件和当前 Excel/WPS 工作表展开。")
    table(
        doc,
        ["模式", "什么时候使用", "系统做什么", "结果"],
        [
            ["模式一：乘系数", "定额已录入，但一批工程数量或编号表达式需要统一乘除系数。", "识别选中定额或章节范围，批量追加或计算系数。", "定额输入表中的工程数量或定额编号表达式被批量更新。"],
            ["模式二：Excel 联动", "工程数量还会变化，需要和 Excel 单元格保持对应。", "保存定额行与 Excel 单元格/表达式的绑定，并可同步。", "后续点击同步即可更新软件中的工程数量。"],
            ["模式三：推荐定额", "拿到工程数量 Excel，需要批量生成定额候选。", "读取 Excel 行，先查组件框，再查定额索引，必要时 AI 补推。", "生成可勾选的推荐表，复制后粘贴到定额输入表。"],
            ["模式四：AI 配对框学习", "某个工程量对应多条定额/材料，希望下次自动推荐这一组。", "读取选中定额和 Excel 选区，AI 配对后写入组件框。", "mapping-boxes 学习库新增可复用组件框。"],
        ],
        [1600, 2500, 2860, 2400],
    )

    h1(doc, "4. 模式一：乘系数模式")
    h2(doc, "4.1 什么时候使用")
    bullet(
        doc,
        [
            "已经录入了一批定额，但设计要求整体调整工程量。",
            "同一章节下所有定额都需要统一乘以折减或放大系数。",
            "需要把系数保留在定额编号表达式中，方便后续查看调整来源。",
        ],
    )
    h2(doc, "4.2 系统做什么")
    body(doc, "系统在定额输入表和章节树右键菜单中加入“乘系数”。用户输入系数后，插件按选中范围批量处理。处理目标可以是“原来的工程量”，也可以是“定额编号”。")
    h2(doc, "4.3 举例说明")
    callout(
        doc,
        "例 1",
        "选中三条混凝土定额，输入 *1.05，选择“乘到原来的工程量”。原工程数量 100、80、20 会变为 105、84、21，工程数量输入列保留计算后的表达或结果。",
    )
    callout(
        doc,
        "例 2",
        "选中一批定额，输入 *0.8，选择“乘到定额编号”。定额编号会形成带系数的表达，适合希望保留调整痕迹的场景。",
    )
    h2(doc, "4.4 操作步骤")
    numbered(
        doc,
        [
            "在定额输入表中选中需要调整的定额行，或在左侧章节树选中目标章节。",
            "右键点击“乘系数”。",
            "选择“乘到原来的工程量”或“乘到定额编号”。",
            "输入系数，例如 *1.05、*0.8 或 /1.1。",
            "确认后系统批量处理，并刷新当前定额输入表。",
        ],
    )
    h2(doc, "4.5 注意事项")
    bullet(
        doc,
        [
            "输入系数前应确认选中范围，章节树操作会影响该章节下多条定额。",
            "乘到工程量会更新工程数量输入和工程数量，适合实际数量调整。",
            "乘到定额编号更偏表达留痕，适合需要在编号侧保留调整系数的场景。",
        ],
    )

    h1(doc, "5. 模式二：Excel 工程量联动模式")
    h2(doc, "5.1 什么时候使用")
    body(doc, "当 Excel 工程数量表是后续变更的主要来源，而软件中定额工程数量需要随着 Excel 更新时，应使用 Excel 工程量联动。")
    h2(doc, "5.2 系统做什么")
    bullet(
        doc,
        [
            "读取 Excel/WPS 当前单元格或选区。",
            "把定额行绑定到单元格地址，例如 E4。",
            "支持表达式绑定，例如 E4+E5、E4*0.8、E4/E5。",
            "支持 AI 智能匹配，让系统根据定额名称、单位、数量和 Excel 行文本建议绑定关系。",
            "同步时读取 Excel 当前值，更新定额输入表工程数量。",
        ],
    )
    h2(doc, "5.3 举例说明")
    callout(
        doc,
        "单元格绑定",
        "定额“C30 混凝土”当前数量来自 Excel 的 E12。用户选中该定额行，在 Excel 点中 E12，点击“绑定 Excel 工程量”。后续 E12 从 100 改为 120 后，点击同步即可把软件工程数量更新为 120。",
    )
    callout(
        doc,
        "表达式绑定",
        "某条定额数量由两个 Excel 单元格相加得到，例如 E12+E13。用户可使用表达式绑定保存 E12+E13，系统同步时自动计算表达式结果。",
    )
    h2(doc, "5.4 操作步骤")
    numbered(
        doc,
        [
            "打开工程数量 Excel/WPS 表。",
            "在软件定额输入表中选中需要绑定的定额行。",
            "在 Excel 中点击对应工程数量单元格，或框选多个单元格。",
            "回到软件右键点击“绑定 Excel 工程量”。",
            "如需连续绑定，选择起始单元格和绑定条数；如需公式，选择表达式绑定。",
            "打开“Excel 联动面板”查看绑定结果。",
            "工程数量变化后，在联动面板点击同步。",
        ],
    )
    h2(doc, "5.5 处理结果")
    body(doc, "同步成功后，软件定额输入表中的工程数量输入和工程数量会更新；联动面板显示最近同步值、状态和更新时间；同步日志记录旧值、新值、定额编号和 Excel 来源。")

    h1(doc, "6. 模式三：推荐定额模式")
    h2(doc, "6.1 什么时候使用")
    body(doc, "当用户拿到一份工程数量 Excel，需要批量找到对应定额时使用推荐定额。它适合初次编制、补充工程量录入、已知工程量但未套定额等场景。")
    h2(doc, "6.2 系统做什么")
    numbered(
        doc,
        [
            "读取 Excel/WPS 当前框选区域；读取失败时可改用剪贴板。",
            "逐行识别工程量名称、单位、数量和上下文。",
            "优先查 mapping-boxes 组件框。",
            "未命中组件框时查 quota-index 全量定额索引。",
            "本地候选不足时由 AI 补推，但不让不可靠结果硬塞。",
            "生成推荐表，默认勾选置信度不低于阈值的行。",
            "用户点击“复制勾选内容”，再粘贴回定额输入表。",
        ],
    )
    h2(doc, "6.3 举例说明")
    table(
        doc,
        ["Excel 工程量", "系统推荐", "说明"],
        [
            ["警示带，m，500", "PY-738 警示（示踪）带铺设", "名称连续词命中，定额名称证据强。"],
            ["土方外运，m3，1200", "LY-21 + LY-34 + LY-35", "如果组件框已学习，则一次返回多条定额。"],
            ["混凝土检查井，座，3", "ZY-41 + 商品混凝土材料", "材料只能通过组件框出现，普通索引不直接返回材料。"],
            ["HPB300钢筋，t，8", "钢筋制作类定额或空推荐", "避免误命中钢筋混凝土构件；不可靠时留空等待扶正。"],
        ],
        [2300, 3000, 4060],
    )
    h2(doc, "6.4 操作步骤")
    numbered(
        doc,
        [
            "在 Excel/WPS 中框选工程量名称、单位和数量所在区域。",
            "回到定额输入表，右键点击“推荐定额”。",
            "查看推荐窗口中的工程量、定额编号、定额名称、单位和换算后数量。",
            "取消勾选不需要复制的推荐，保留正确推荐。",
            "点击“复制勾选内容”。",
            "回到定额输入表目标位置，从定额编号列开始粘贴。",
            "发现错误推荐时，选中正确定额并点击对应工程量行的“扶正”。",
        ],
    )
    h2(doc, "6.5 数量换算")
    body(doc, "系统复制时会按 Excel 单位和定额单位做数量级换算。例：Excel 单位为 m3、数量为 1200，定额单位为 100m3，则复制数量为 12。单位不能可靠换算时，保留原数量或等待人工确认。")

    h1(doc, "7. 模式四：AI 配对框学习模式")
    h2(doc, "7.1 什么时候使用")
    body(doc, "当一个工程量经常对应多条定额或材料时，普通关键词检索往往不够。例如“土方外运”可能对应挖装、运输和增运；“混凝土检查井”可能对应定额和商品混凝土材料。此时应使用 AI 配对框学习。")
    h2(doc, "7.2 系统做什么")
    body(doc, "系统读取定额输入表中已选中的定额/材料，再读取 Excel/WPS 中框选的工程量行。它先用本地数量关系做预匹配，再用 AI 辅助判断名称和上下文，最后生成可确认的配对候选。用户确认后，组件框写入 mapping-boxes.jsonl。")
    h2(doc, "7.3 举例说明")
    callout(
        doc,
        "土方外运组件框",
        "用户在定额输入表选中 LY-21、LY-34、LY-35，在 Excel 中框选“土方外运 m3 1200”。AI 配对后保存组件框。下次推荐遇到土方外运时，系统直接返回 LY-21、LY-34、LY-35 三行。",
    )
    h2(doc, "7.4 操作步骤")
    numbered(
        doc,
        [
            "在定额输入表中选中一条或多条正确的定额/材料。",
            "在 Excel/WPS 中框选对应工程量行。",
            "右键点击“添加对应框内容”。",
            "点击“AI 配对”。",
            "查看配对结果，确认工程量名称、单位、数量和目标定额/材料。",
            "勾选正确配对并确认写入。",
            "重新打开推荐定额窗口验证组件框是否优先命中。",
        ],
    )
    h2(doc, "7.5 与批量反向建模板的关系")
    body(doc, "AI 配对框学习和批量反向建模板的本质相同，都是把“工程量表达 -> 一组定额/材料”写入 mapping-boxes。当前功能适合小范围人工确认，已经能覆盖主要需求，因此暂不新增批量反向建模板，避免功能重复。")

    h1(doc, "8. 用户映射学习功能")
    h2(doc, "8.1 什么是用户映射")
    body(doc, "用户映射是指用户人工确认过的工程量与定额/材料之间的对应关系。它保存在 mapping-boxes.jsonl 中，是推荐链路中优先级最高的数据来源。")
    h2(doc, "8.2 工作原理")
    numbered(
        doc,
        [
            "用户通过扶正或 AI 配对确认正确关系。",
            "系统记录工程量名称、单位、样本权重、目标定额/材料、使用时间等信息。",
            "下次读取 Excel 工程量时，系统先查 mapping-boxes。",
            "如果工程量表达相同或相似，直接返回组件框。",
            "用户复制勾选内容视为接受推荐，样本权重提高。",
        ],
    )
    h2(doc, "8.3 影响范围")
    bullet(
        doc,
        [
            "影响推荐定额：组件框命中优先于普通索引检索。",
            "影响材料返回：材料只能通过组件框出现。",
            "不影响原软件业务数据库：推荐定额仍通过剪贴板粘贴。",
            "不覆盖人工扶正：旧学习样本不能覆盖人工扶正框。",
        ],
    )

    h1(doc, "9. 典型工作场景")
    h2(doc, "9.1 场景一：新工程量表初次推荐定额")
    body(doc, "背景：设计单位提供了一份新的工程数量表，用户希望快速生成定额输入候选。")
    numbered(
        doc,
        [
            "在 Excel 中框选工程量名称、单位、数量三列。",
            "在定额输入表右键打开推荐定额。",
            "检查系统返回的定额编号、名称、单位和换算数量。",
            "对高置信度结果直接复制粘贴。",
            "对错误或空推荐，使用扶正或 AI 配对框学习。",
        ],
    )
    body(doc, "预期效果：大多数常见工程量可以直接推荐；复杂工程量通过组件框逐步沉淀，后续准确率提高。")

    h2(doc, "9.2 场景二：设计变更后更新数量")
    body(doc, "背景：Excel 工程数量表发生变更，已录入软件的定额数量需要同步。")
    numbered(
        doc,
        [
            "前期把关键定额行绑定到 Excel 单元格或表达式。",
            "设计变更后更新 Excel。",
            "打开 Excel 联动面板。",
            "点击同步，系统读取最新单元格值并更新定额输入表。",
            "查看同步日志，核对旧值和新值。",
        ],
    )

    h2(doc, "9.3 场景三：一项工程量对应多条定额")
    body(doc, "背景：普通推荐只能返回一条最优定额，但现场某些工程量需要一组定额/材料。")
    numbered(
        doc,
        [
            "在定额输入表中选中这组正确的定额/材料。",
            "在 Excel 中框选对应工程量。",
            "使用“添加对应框内容”进行 AI 配对。",
            "确认写入 mapping-boxes。",
            "后续推荐时，该工程量会返回多行组件推荐。",
        ],
    )

    h2(doc, "9.4 场景四：章节内批量调整系数")
    body(doc, "背景：某个章节内已录入定额需要统一乘以调整系数。")
    numbered(
        doc,
        [
            "在左侧章节树选中目标章节。",
            "右键点击“乘系数”。",
            "输入系数并选择作用目标。",
            "系统批量更新该章节下定额。",
            "刷新后抽查关键定额工程数量。",
        ],
    )

    h1(doc, "10. 常见问题解答")
    table(
        doc,
        ["问题", "回答"],
        [
            ["Q1：推荐定额会直接写项目数据库吗？", "不会。推荐结果通过剪贴板粘贴回原软件，保持原软件业务逻辑。"],
            ["Q2：为什么普通推荐不返回材料？", "材料必须通过人工扶正后的组件框出现，避免普通关键词把材料误当定额推荐。"],
            ["Q3：推荐错了怎么办？", "在软件中选中正确定额，点击推荐行的“扶正”；或者使用“添加对应框内容”保存组件框。"],
            ["Q4：Excel 联动是不是自动实时同步？", "当前以手动同步和面板同步为主，避免 Excel 临时编辑导致软件数据被误改。"],
            ["Q5：AI 没配置还能用吗？", "可以。乘系数、Excel 基础绑定、推荐定额本地索引和人工扶正仍可使用；AI 只增强配对和补推。"],
            ["Q6：mapping-boxes 满了怎么办？", "每个框默认最多保留 30 个工程量样本，容量满时应优先淘汰低权重且长期未使用样本。"],
            ["Q7：为什么有些推荐为空？", "系统宁可留空也不强塞低可靠结果。空推荐行仍保留工程量名称、单位、数量和扶正按钮。"],
        ],
        [3100, 6260],
    )

    h1(doc, "11. 操作技巧和建议")
    h2(doc, "11.1 提高推荐准确率")
    bullet(
        doc,
        [
            "Excel 工程量名称尽量完整，不要只保留单字或过短简称。",
            "单位和数量列尽量一起框选，便于单位换算和数量校验。",
            "遇到一对多工程量，优先使用 AI 配对框保存组件框，而不是反复手动粘贴。",
            "对钢筋、检查井、土方外运等高频复杂项，先建立稳定组件框。",
        ],
    )
    h2(doc, "11.2 提高 Excel 联动可靠性")
    bullet(
        doc,
        [
            "绑定前确认 Excel/WPS 文件已保存，工作表名称稳定。",
            "表达式尽量简单，优先使用单元格加减乘除，不使用跨工作簿复杂公式。",
            "同步前先确认 Excel 不处于单元格编辑状态。",
            "大范围同步后抽查关键定额数量。",
        ],
    )
    h2(doc, "11.3 质量控制要点")
    bullet(
        doc,
        [
            "组件框命中时只第一行显示工程量名称和扶正按钮，后续行留空，这是正常表现。",
            "数量换算后应核对单位，例如 m3 与 100m3、m2 与 100m2。",
            "普通推荐只保留一条最优定额，出现多条需求时应通过组件框解决。",
            "不要删除 learning.jsonl、mapping-boxes.jsonl、quota-index.jsonl 等本地数据文件。",
        ],
    )

    h1(doc, "12. 附录：数据文件和术语说明")
    h2(doc, "12.1 主要数据文件")
    table(
        doc,
        ["文件", "用途", "说明"],
        [
            ["quota-index.jsonl", "预算定额索引", "普通推荐的主要检索来源。"],
            ["material-index.jsonl", "材料索引", "仅供组件框和材料目标补全使用。"],
            ["mapping-boxes.jsonl", "组件框学习库", "保存人工确认后的工程量表达与一组定额/材料关系。"],
            ["learning.jsonl", "旧学习样本", "保留兼容和备份，不应删除。"],
            ["Excel 联动 XML", "项目绑定关系", "保存定额行与 Excel 单元格/表达式的关系。"],
        ],
        [2500, 2500, 4360],
    )
    h2(doc, "12.2 术语说明")
    table(
        doc,
        ["术语", "说明"],
        [
            ["组件框", "一个工程量表达对应一条或多条定额/材料的组合。"],
            ["扶正", "用户把错误推荐纠正为正确定额/组件框的操作。"],
            ["AI 配对", "系统根据选中定额和 Excel 工程量生成配对候选，用户确认后保存。"],
            ["工程数量输入", "软件定额输入表中用户可编辑或粘贴的工程数量字段。"],
            ["换算数量", "根据 Excel 单位和定额单位换算后的定额工程数量。"],
        ],
        [2400, 6960],
    )

    h2(doc, "12.3 与参考 PDF 标题结构对应关系")
    table(
        doc,
        ["参考 PDF 标题", "本文档对应内容"],
        [
            ["系统简介", "1.1 系统简介"],
            ["核心价值", "1.2 核心价值"],
            ["四种处理模式", "3. 四种处理模式"],
            ["模式一/二/三/四", "4-7 章分别说明四个插件模式"],
            ["用户映射学习功能", "8. 用户映射学习功能"],
            ["典型工作场景", "9. 典型工作场景"],
            ["常见问题解答", "10. 常见问题解答"],
            ["操作技巧和建议", "11. 操作技巧和建议"],
            ["附录", "12. 附录：数据文件和术语说明"],
        ],
        [3000, 6360],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
